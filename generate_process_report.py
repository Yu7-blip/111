#!/usr/bin/env python3
"""生成过程报告 + 分工明细表 — 独立分工版"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement
import datetime

def setup_doc():
    doc = Document()
    s = doc.styles['Normal']
    s.font.name = '宋体'; s.font.size = Pt(11); s.font.color.rgb = RGBColor(0,0,0)
    s.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    s.paragraph_format.line_spacing = 1.5; s.paragraph_format.space_after = Pt(4)
    for i in range(1,4):
        hs = doc.styles[f'Heading {i}']
        hs.font.color.rgb = RGBColor(0,0,0); hs.font.name = '黑体'
        hs.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        if i==1: hs.font.size=Pt(18); hs.font.bold=True
        elif i==2: hs.font.size=Pt(14); hs.font.bold=True
        else: hs.font.size=Pt(12); hs.font.bold=True
    for sec in doc.sections:
        sec.top_margin=Cm(2.54); sec.bottom_margin=Cm(2.54)
        sec.left_margin=Cm(3.17); sec.right_margin=Cm(3.17)
    return doc

def P(doc, text, bold=False, indent=True, size=11, align=None):
    p = doc.add_paragraph()
    if indent: p.paragraph_format.first_line_indent = Pt(22)
    p.paragraph_format.line_spacing = 1.5
    if align: p.alignment = align
    r = p.add_run(text); r.bold=bold; r.font.size=Pt(size); r.font.color.rgb=RGBColor(0,0,0)
    r.font.name='宋体'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')
    return p

def B(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent=Cm(1.0); p.paragraph_format.first_line_indent=Cm(-0.4)
    p.paragraph_format.line_spacing=1.5
    r = p.add_run('• '+text); r.font.size=Pt(10.5); r.font.color.rgb=RGBColor(0,0,0)
    r.font.name='宋体'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')
    return p

def CODE(doc, text):
    t = doc.add_table(rows=1, cols=1); t.alignment=WD_TABLE_ALIGNMENT.CENTER
    c = t.rows[0].cells[0]
    c._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5" w:val="clear"/>'))
    tc_pr=c._tc.get_or_add_tcPr(); tb=OxmlElement('w:tcBorders')
    for e in ('top','left','bottom','right'):
        bd=OxmlElement(f'w:{e}'); bd.set(qn('w:val'),'single'); bd.set(qn('w:sz'),'4')
        bd.set(qn('w:space'),'0'); bd.set(qn('w:color'),'999999'); tb.append(bd)
    tc_pr.append(tb); c.paragraphs[0].clear()
    for i,ln in enumerate(text.split('\n')):
        p = c.paragraphs[0] if i==0 else c.add_paragraph()
        p.paragraph_format.space_before=Pt(1); p.paragraph_format.space_after=Pt(1)
        p.paragraph_format.line_spacing=1.2
        r=p.add_run(ln); r.font.name='Consolas'; r.font.size=Pt(8.5); r.font.color.rgb=RGBColor(51,51,51)
    doc.add_paragraph()

def TBL(doc, hd, rows):
    t=doc.add_table(rows=1+len(rows), cols=len(hd)); t.style='Table Grid'
    t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=True
    for i,h in enumerate(hd):
        c=t.rows[0].cells[i]; c.text=''; p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        r=p.add_run(h); r.bold=True; r.font.size=Pt(9); r.font.name='黑体'
        r._element.rPr.rFonts.set(qn('w:eastAsia'),'黑体'); r.font.color.rgb=RGBColor(0,0,0)
        c._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="E0E0E0" w:val="clear"/>'))
        c.vertical_alignment=WD_ALIGN_VERTICAL.CENTER
    for ri,row in enumerate(rows):
        for ci,val in enumerate(row):
            c=t.rows[ri+1].cells[ci]; c.text=''; p=c.paragraphs[0]
            r=p.add_run(str(val)); r.font.size=Pt(9); r.font.name='宋体'
            r._element.rPr.rFonts.set(qn('w:eastAsia'),'宋体'); r.font.color.rgb=RGBColor(0,0,0)
            c.vertical_alignment=WD_ALIGN_VERTICAL.CENTER
    doc.add_paragraph()

def SEP(doc):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(2); p.paragraph_format.space_after=Pt(2)
    pPr=p._p.get_or_add_pPr(); pBdr=OxmlElement('w:pBdr')
    b=OxmlElement('w:bottom'); b.set(qn('w:val'),'single'); b.set(qn('w:sz'),'4')
    b.set(qn('w:space'),'0'); b.set(qn('w:color'),'CCCCCC'); pBdr.append(b); pPr.append(pBdr)

def COVER(doc, title, subtitle, lines):
    for _ in range(5): doc.add_paragraph()
    tp=doc.add_paragraph(); tp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    tr=tp.add_run(title); tr.font.size=Pt(36); tr.bold=True
    tr.font.color.rgb=RGBColor(0,0,0); tr.font.name='黑体'
    tr._element.rPr.rFonts.set(qn('w:eastAsia'),'黑体')
    doc.add_paragraph()
    sp=doc.add_paragraph(); sp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    sr=sp.add_run(subtitle); sr.font.size=Pt(22)
    sr.font.color.rgb=RGBColor(0,0,0); sr.font.name='黑体'
    sr._element.rPr.rFonts.set(qn('w:eastAsia'),'黑体')
    for _ in range(4): doc.add_paragraph()
    for ln in lines:
        ip=doc.add_paragraph(); ip.alignment=WD_ALIGN_PARAGRAPH.CENTER
        ir=ip.add_run(ln); ir.font.size=Pt(13); ir.font.color.rgb=RGBColor(0,0,0)
        ir.font.name='宋体'; ir._element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')
    doc.add_page_break()


# ============================================================
#  过程报告
# ============================================================
def build_process_report():
    doc = setup_doc()
    COVER(doc, '外卖配送平台', '项目开发过程报告', [
        '组  别：第 11 组',
        '所属课程：软件工程课程设计',
        f'日  期：{datetime.date.today().strftime("%Y年%m月%d日")}',
    ])

    # ---- 一、团队信息 ----
    doc.add_heading('一、团队信息', level=1)
    TBL(doc,
        ['学号', '姓名', '班级', '组内角色', '负责模块'],
        [
            ['2300770172', '余  睿', '软工2304', '组长',
             '系统架构 + 订单模块 + 派单算法 + GeoHash\n+ 缓存防护 + 安全体系 + Admin营销/系统页面'],
            ['2300770178', '赵宇弦', '软工2304', '组员',
             '认证模块 + 商品模块 + 商家搜索\n+ Admin/Merchant前端 + 小程序C端页面'],
            ['2300770046', '周峥嵘', '软工2304', '组员',
             '配送模块 + Outbox/事件系统 + 评价/反馈\n+ Admin用户/商家/骑手页面 + 数据库 + 测试'],
        ]
    )
    P(doc, '本组三人各自负责独立模块，明确分工边界，通过Git版本管理确保代码不冲突。')

    # ---- 二、选题与背景 ----
    doc.add_heading('二、项目选题与背景', level=1)
    P(doc, '本组选择"外卖配送平台"作为课程设计题目。外卖O2O是当前互联网行业最成熟的业务场景之一，'
      '其技术栈覆盖GeoHash空间索引、实时位置追踪、智能调度算法、分布式事务一致性以及高并发缓存优化等'
      '多个软件工程核心领域，非常适合作为课程设计的综合实践项目。通过完整开发一款覆盖用户端、商家端、'
      '骑手端和后台管理四大角色的外卖平台，可以系统性地锻炼前后端分离开发、数据库设计和中间件集成等关键能力。')

    # ---- 三、开发计划与进度 ----
    doc.add_heading('三、开发计划与进度安排', level=1)
    P(doc, '项目总周期16周，划分为五个阶段，各阶段三人并行推进各自负责的模块。')
    TBL(doc,
        ['阶段', '时间', '主要任务', '产出'],
        [
            ['需求分析', '第1-2周', '调研竞品、明确49项功能需求\n编写需求规格说明书', '需求文档 + 功能清单'],
            ['系统设计', '第3-5周', '架构设计、数据库建模\n接口定义、前端原型设计', 'ER图 + 架构图 + API文档'],
            ['编码实现', '第6-11周', '三人各自独立开发所负责的后端模块\n和前端页面，按周推进', '可运行版本 (MVP)'],
            ['集成测试', '第12-14周', '前后端联调、功能测试\nBug修复、回归验证', '测试报告'],
            ['文档答辩', '第15-16周', '撰写综合设计报告和过程报告\n准备答辩PPT', '最终报告 + 答辩材料'],
        ]
    )

    # ---- 四/五/六 — 需求/设计/编码 ----
    doc.add_heading('四、需求分析过程', level=1)
    P(doc, '三人共同对美团和饿了么进行功能拆解后，将需求整理为三大模块49项子需求。'
      '余睿负责梳理模块1（用户端与商家端14项）和模块3（后台管理15项），'
      '赵宇弦负责模块2（骑手端与配送系统14项），周峥嵘负责6项非功能性需求整理。'
      '最终汇总形成完整的需求规格文档。')

    doc.add_heading('五、系统设计过程', level=1)
    P(doc, '系统设计阶段，余睿负责整体架构设计——确定Spring Boot 3.2.5 + MyBatis-Plus技术选型、'
      '四层分层架构（过滤器→拦截器→Controller→Service→Mapper）、以及安全防护体系方案。'
      '周峥嵘负责数据库设计——完成20张表的E-R建模、字段定义和索引策略设计。'
      '赵宇弦负责前端架构——确定Vue 3 + Element Plus技术栈、路由设计和组件结构。'
      '各自完成设计后分别输出对应文档。')

    doc.add_heading('六、编码实现过程', level=1)

    P(doc, '6.1 项目骨架搭建（第3-5周）', bold=True)
    P(doc, '余睿搭建Spring Boot后端项目骨架，完成pom.xml依赖配置、application.yml多环境配置、'
      '以及Redis/WebSocket/Swagger/CORS等9个配置类。赵宇弦搭建Admin和Merchant两个Vue 3前端项目，'
      '配置Vite代理、路由和Axios拦截器。周峥嵘编写schema.sql（384行）和数据种子data.sql（88行），'
      '三人各自在本地完成开发环境部署。')

    P(doc, '6.2 余睿 — 订单模块与核心基础设施（第6-9周）', bold=True)
    P(doc, '第6-7周：实现认证拦截器链——AdminLoginInterceptor、MerchantLoginInterceptor、'
      'WxLoginInterceptor三套JWT拦截器，以及XssFilter、SecurityHeadersFilter安全过滤器。'
      '实现RateLimitInterceptor接口限流和Resilience4j熔断降级配置。')
    P(doc, '第7-8周：自研GeoHashUtil（158行），实现Base32编码/解码/邻域搜索/精度自适应四个核心算法。'
      '实现ShopServiceImpl（495行）中商家入驻审核和GeoHash附近搜索功能。'
      '封装TencentMapService（300行）集成腾讯地图Geocode和骑行路径API。')
    P(doc, '第8-9周：实现OrderServiceImpl（893行），包含订单状态机（8种状态）、下单优惠自动计算'
      '（优惠券+满减活动叠加）、支付沙箱模拟（90%成功率）、退款三级仲裁、大订单拆单等完整业务逻辑。')
    P(doc, '第9周：实现DispatchServiceImpl（354行）四维加权智能派单算法，'
      '设计三级分流策略（自动派单≥0.60 / 推送抢单≥0.35 / 自由抢单）。'
      '搭建WebSocket实时推送（OrderNotificationEndpoint + NotificationService）。')

    P(doc, '6.3 余睿 — 缓存与安全加固（第10-11周）', bold=True)
    P(doc, '实现CacheUtil（183行）缓存三防体系——Redis Bitmap布隆过滤器防穿透、'
      'ConcurrentHashMap互斥锁防击穿、随机TTL±30%防雪崩。'
      '编写Admin后台营销配置页面（CouponList 139行 + FullReduce 232行）和系统配置页面（ConfigList 172行）。')

    P(doc, '6.4 赵宇弦 — 认证、商品模块与前端开发（第6-9周）', bold=True)
    P(doc, '第6-7周：实现AuthServiceImpl（311行）用户认证模块——手机号+密码注册登录、'
      'BCrypt密码加密、JwtUtil Token签发。实现GoodsServiceImpl（280行）商品CRUD、分类管理、'
      '上下架和库存控制。')
    P(doc, '第7-9周：完成Admin管理后台全部页面的API层封装（12个模块）及数据看板（ECharts图表）、'
      '订单管理（344行）、评价管理页面。搭建Merchant商家管理后台全部6个页面，集成@wangeditor富文本编辑器。')
    P(doc, '第9-10周：开发微信小程序C端全部页面——首页（162行）、商家列表（68行）、菜单浏览（299行）、'
      '下单确认（182行）、支付页（142行）、订单列表（146行）、订单详情（133行）、优惠券中心（60行）、'
      '个人中心（116行），以及7个公共组件。')

    P(doc, '6.5 周峥嵘 — 配送系统与事件系统（第8-10周）', bold=True)
    P(doc, '第8-9周：实现DeliveryServiceImpl（877行）骑手端全部业务——注册与实名认证、接单/拒单、'
      'GPS位置实时上报（Redis setex + 定时批量落DB）、收入统计与提现申请。'
      '实现WxDeliveryController（310行）骑手相关全部REST接口。')
    P(doc, '第9-10周：实现EventLogService（184行）Transaction Outbox模式——saveEvent事务写入、'
      'tryPublishAfterCommit实时推送、processStaleEvents定时兜底扫描（LIMIT 50分批，最多重试5次）。'
      '实现OrderEventListener（207行）异步事件监听器，处理ORDER_PAID等5种事件类型。')
    P(doc, '第10周：实现OrderScheduledTask（192行）5个定时任务——自动取消未支付订单（15分钟）、'
      '配送超时处理（60分钟）、优惠券自动过期、满减活动自动启停、事件表兜底扫描。')

    P(doc, '6.6 周峥嵘 — 后台管理与小程序骑手端（第11周）', bold=True)
    P(doc, '实现评价系统WxEvaluationController（162行），含评价提交、骑手好评率和等级自动重算、'
      '店铺评分更新。实现FeedbackController（105行）反馈/投诉/申诉系统。'
      '开发Admin后台用户管理页面（UserList 208行）、商家审核页面（MerchantAudit 136行）、'
      '骑手管理页面（DeliveryList 227行）。开发小程序骑手端页面——配送中（delivering 490行，'
      '含腾讯地图SDK集成）、收入统计（62行）、骑手资料（138行），以及地址管理（304行）和反馈页面（33行）。')

    # ---- 七、测试过程 ----
    doc.add_heading('七、测试过程', level=1)
    P(doc, '测试阶段由周峥嵘主要负责。编写了15个功能测试用例覆盖全链路场景，逐一执行并记录Bug。'
      '余睿负责修复订单模块和派单模块发现的问题，赵宇弦负责修复前端页面展示和接口联调问题。'
      '全组使用Knife4j在线文档对27个Controller的REST接口进行调试验证，'
      '经三轮回归测试后核心业务流程全部通过。')

    # ---- 八、遇到的问题与解决方案 ----
    doc.add_heading('八、开发中遇到的问题与解决方案', level=1)
    issues = [
        ('问题1：GeoHash搜索精度与性能的平衡（余睿）',
         '初期将精度固定为7级，大半径搜索时遗漏远处商家。解决方案：实现precisionForRadius方法，'
         '根据搜索半径自适应选择精度等级，始终搜索9个单元格（中心+8邻域），再用Haversine精确过滤。'),
        ('问题2：订单支付后库存更新的一致性（余睿）',
         '最初同步更新库存导致接口响应慢且缺乏失败补偿。引入Transaction Outbox模式——'
         '事件日志与业务在同一事务中写入，事务提交后异步处理库存变更，Redis失败则由定时任务兜底。'),
        ('问题3：派单算法中顺路度的量化（余睿）',
         '设计两维度判断——目的地距离<1km（+0.3分）+ 方位角差异<30°（+0.2分），有效量化"顺路"概念。'),
        ('问题4：WebSocket在Vite代理下的连接失败（赵宇弦）',
         '开发环境中WebSocket频繁断开。排查发现Vite dev server代理未转发WebSocket升级请求，'
         '在vite.config.js中添加ws:true配置后解决。'),
        ('问题5：GPS高频写入的存储压力（周峥嵘）',
         '骑手每5秒上报一次GPS，直接写DB压力大。采用Redis先缓存（setex 30秒过期）+ 定时批量落DB策略，'
         '兼顾实时查询和存储效率。'),
    ]
    for t,d in issues:
        P(doc, t, bold=True); P(doc, d)

    # ---- 九、个人总结 ----
    doc.add_heading('九、个人总结与收获', level=1)

    P(doc, '9.1 余睿（组长，2300770172）', bold=True)
    P(doc, '作为组长，我负责系统架构设计和后端最核心的几个模块——订单系统（893行）、智能派单算法（354行）、'
      '自研GeoHash（158行）、缓存三防体系（183行）以及安全防护体系。技术层面最大的收获是分布式事务的实践'
      '——通过Transaction Outbox模式将"最终一致性"这个抽象概念变成了具体可运行的代码。'
      'GeoHash算法的从零实现也让我深刻理解了空间索引在LBS场景中的工程价值。派单算法的四维加权设计则锻炼了'
      '将业务需求转化为数学模型的能力。在前端方面我也独立完成了Admin后台的营销配置和系统管理页面。'
      '同时作为组长，我负责统筹项目进度、制定Git分支规范和最终报告的统稿工作，'
      '这些管理经验与技术实践一样宝贵。')

    P(doc, '9.2 赵宇弦（组员，2300770178）', bold=True)
    P(doc, '我负责用户认证模块（311行）、商品管理模块（280行）的后端开发，以及全部前端项目的开发——'
      '包括Admin管理后台的架构搭建和7个核心页面（数据看板、订单管理、评价管理等）、'
      'Merchant商家管理后台的全部6个页面（集成@wangeditor富文本编辑器）、'
      '以及微信小程序C端用户的全部10个页面和7个公共组件，总计约6,700行前端代码。'
      '通过从零搭建两个Vue 3管理后台和一个完整的小程序端，我深入掌握了Vue 3 Composition API、'
      'Element Plus组件库深度定制、ECharts数据可视化以及微信小程序原生开发的全套技能。'
      '后端开发也让我对Spring Boot的拦截器机制和JWT鉴权有了实际应用经验。')

    P(doc, '9.3 周峥嵘（组员，2300770046）', bold=True)
    P(doc, '我负责数据库设计（20张表DDL + 索引策略 + 种子数据）、配送系统后端开发（877行）、'
      'Transaction Outbox事件系统、评价与反馈模块、5个定时任务，以及Admin后台用户管理、商家审核、'
      '骑手管理三个前端页面和小程序骑手端6个页面（含配送中页面的腾讯地图SDK集成），总计约1,800行后端代码'
      '和1,200行前端代码。测试阶段我作为主测试人编写了15个测试用例并逐项验证。'
      '通过这次课设，我从只会写SQL到能够独立设计数据库模型，从只接触过后端到能开发Vue前端页面和小程序页面，'
      '还深入理解了Transaction Outbox分布式事务模式和定时任务调度机制，技能面得到了极大的拓展。')

    SEP(doc)
    P(doc, '')
    P(doc, '本次课程设计历时16周，三位组员各自独立负责明确的模块，通过Git版本管理确保高效并行开发。'
      '项目代码总量超过30,000行，前后端接口全部联调通过，是一个完整可交付的软件系统。')

    doc.save('D:/SpringbootFinal/_temp_process_report.docx')
    print('Process report saved.')


# ============================================================
#  分工明细表
# ============================================================
def build_division_table():
    doc = setup_doc()
    COVER(doc, '外卖配送平台', '团队分工明细表', [
        '组  别：第 11 组',
        '所属课程：软件工程课程设计',
    ])

    # ---- 一、基本信息 ----
    doc.add_heading('一、团队成员基本信息', level=1)
    TBL(doc,
        ['学号', '姓名', '班级', '组内角色', '负责方向概述'],
        [
            ['2300770172', '余  睿', '软工2304', '组长',
             '系统架构设计 + 订单模块 + 派单算法\n+ GeoHash + 缓存防护 + 安全体系\n+ Admin营销/系统配置页面'],
            ['2300770178', '赵宇弦', '软工2304', '组员',
             '认证模块 + 商品模块 + 商家搜索\n+ Admin前端 + Merchant前端\n+ 小程序C端全部页面'],
            ['2300770046', '周峥嵘', '软工2304', '组员',
             '配送模块 + Outbox事件 + 评价/反馈\n+ 定时任务 + 数据库设计\n+ Admin用户/商家/骑手页面 + 小程序骑手端'],
        ]
    )

    # ---- 二、后端分工 ----
    doc.add_heading('二、后端开发分工', level=1)
    P(doc, '说明：后端共116个Java类，三大模块由三人分别独立负责，模块之间通过REST API交互。')
    TBL(doc,
        ['模块', '主要文件', '负责人', '工作量'],
        [
            ['项目骨架\n+ 配置层', 'BackendApplication\napplication*.yml\nconfig包(9个类)', '余  睿', '约500行'],
            ['认证模块', 'AuthServiceImpl(311行)\nWxAuthController\nJwtUtil(76行)', '赵宇弦', '约400行'],
            ['JWT拦截器\n+ 安全防护', 'AdminLoginInterceptor\nMerchantLoginInterceptor\nWxLoginInterceptor\n'
             'XssFilter(44行)\nSecurityHeadersFilter(22行)\nRateLimitInterceptor(80行)\nCircuitBreakerFallback(82行)',
             '余  睿', '约350行'],
            ['商家/商品\n+ GeoHash', 'ShopServiceImpl(495行)\nGoodsServiceImpl(280行)\nGeoHashUtil(158行)\nGeoUtil(39行)',
             '赵宇弦\n(商品+商家CRUD)\n余  睿\n(GeoHash+搜索)', '约970行'],
            ['订单模块', 'OrderServiceImpl(893行)\nWxOrderController(111行)\nOrder+OrderItem实体',
             '余  睿', '约1100行'],
            ['智能派单', 'DispatchServiceImpl(354行)', '余  睿', '约350行'],
            ['配送系统', 'DeliveryServiceImpl(877行)\nWxDeliveryController(310行)\nDelivery+DeliveryRecord+DeliveryTrack实体',
             '周峥嵘', '约1300行'],
            ['Transaction Outbox\n+ 事件系统', 'EventLogService(184行)\nOrderEventListener(207行)\nEventMessage+RedisMessagePublisher(61行)',
             '周峥嵘', '约450行'],
            ['评价+反馈', 'WxEvaluationController(162行)\nFeedbackController(105行)\nEvaluation+Feedback实体',
             '周峥嵘', '约350行'],
            ['定时任务', 'OrderScheduledTask(192行)\n5个@Scheduled方法', '周峥嵘', '约200行'],
            ['缓存防护', 'CacheUtil(183行)\nRedisUtil(63行)', '余  睿', '约250行'],
            ['WebSocket\n+ 地图服务', 'OrderNotificationEndpoint(64行)\nNotificationService(38行)\nTencentMapService(300行)',
             '余  睿', '约420行'],
            ['Admin后台接口', 'AdminDashboardController\nAdminMerchant/Order/User/Delivery/\nEvaluation/Marketing/System等Controller',
             '余  睿\n(4个Controller)\n周峥嵘\n(5个Controller)', '约800行'],
        ]
    )

    # ---- 三、前端分工 ----
    doc.add_heading('三、前端开发分工', level=1)
    P(doc, '说明：前端三个项目由三人各自独立负责不同页面。')
    TBL(doc,
        ['前端项目', '模块/页面', '负责人', '工作量'],
        [
            ['Admin\n管理后台\n(15页)', '项目初始化 + Layout布局 + 路由 + 状态管理\n+ API层封装(12个模块)\n+ 数据看板(ECharts)\n+ 订单管理\n+ 评价管理',
             '赵宇弦', '约2,200行'],
            ['Admin\n管理后台', '营销配置(优惠券+满减活动)\n+ 系统配置 + 反馈处理',
             '余  睿', '约520行'],
            ['Admin\n管理后台', '用户管理 + 商家审核\n+ 骑手管理',
             '周峥嵘', '约570行'],
            ['Merchant\n商家后台\n(6页)', '全部6个路由页面:\n登录/看板/店铺设置/商品管理\n/订单处理/优惠券管理\n+ 富文本编辑器组件',
             '赵宇弦', '约1,500行'],
            ['微信小程序\n(C端用户)', '首页 + 商家列表 + 菜单浏览\n+ 下单确认 + 支付页\n+ 订单列表 + 订单详情\n+ 优惠券中心 + 个人中心',
             '赵宇弦', '约1,600行'],
            ['微信小程序\n(C端用户)', '收货地址管理(含地图选点)\n+ 反馈/投诉提交',
             '周峥嵘', '约340行'],
            ['微信小程序\n(骑手端)', '接单大厅 + 任务池',
             '赵宇弦', '约250行'],
            ['微信小程序\n(骑手端)', '配送中(含腾讯地图导航,490行)\n+ 收入统计 + 骑手资料\n+ 骑手评价查看',
             '周峥嵘', '约750行'],
            ['小程序\n公共组件', '7个组件:\ncart-popup/shop-card/order-card\n/star-rating 等\n+ utils工具库(5个文件)',
             '赵宇弦', '约800行'],
        ]
    )

    # ---- 四、数据库与文档 ----
    doc.add_heading('四、数据库设计与文档分工', level=1)
    TBL(doc,
        ['任务', '具体内容', '负责人', '工作量'],
        [
            ['数据库设计', 'E-R建模 + 20张表DDL\n+ 索引策略', '周峥嵘', 'schema.sql 384行'],
            ['种子数据', '管理员/用户/商家/商品/优惠券测试数据', '周峥嵘', 'data.sql 88行'],
            ['增量修复脚本', '6个fix_*.sql', '周峥嵘', '约130行'],
            ['接口联调', '27个Controller的Swagger调试', '三人各自负责', '每人约9个'],
            ['功能测试', '15个测试用例编写与执行验证', '周峥嵘', '测试用例文档'],
            ['综合设计报告', '第1/2/4/6章(概述/需求/设计/总结)', '余  睿', '约12,000字'],
            ['综合设计报告', '第3章(数据库设计) + 第5章(测试与运行)', '周峥嵘', '约3,000字'],
            ['过程报告', '全过程记录 + 分工明细表', '余  睿', '约5,000字'],
            ['答辩PPT', 'PPT制作与汇总', '余  睿', '16页'],
        ]
    )

    # ---- 五、工作量统计 ----
    doc.add_heading('五、工作量统计', level=1)
    TBL(doc,
        ['成员', '后端开发', '前端开发', '数据库/测试/文档', '总计'],
        [
            ['余  睿\n(组长)', '订单+派单+GeoHash+缓存\n+安全+WebSocket+地图\n+Admin接口\n≈ 3,000行',
             'Admin后台:\n营销配置+系统配置+反馈\n≈ 520行',
             '系统架构设计\n+ 综合报告主体\n+ 过程报告 + PPT\n+ 项目进度管理',
             '约3,500行代码\n+ 17,000字文档'],
            ['赵宇弦\n(组员)', '认证模块+商品模块\n≈ 800行',
             'Admin前端架构+4页\n+ Merchant全部6页\n+ 小程序C端+部分骑手端\n+ 公共组件\n≈ 6,400行',
             '前端架构搭建\n+ 接口联调',
             '约7,200行代码'],
            ['周峥嵘\n(组员)', '配送+Outbox+评价+反馈\n+定时任务+Admin接口\n≈ 1,800行',
             'Admin用户/商家/骑手页面\n+ 小程序地址+骑手端\n≈ 1,200行',
             '数据库设计(500行SQL)\n+ 15个测试用例\n+ 综合报告2个章节',
             '约3,000行代码\n+ 500行SQL\n+ 3,000字文档'],
        ]
    )

    SEP(doc)
    P(doc, '')
    P(doc, '注：上表代码行数为估算值（含注释和空行）。三人各自独立负责明确模块，'
      '模块间通过REST API接口契约交互，互不依赖内部实现细节。')

    doc.save('D:/SpringbootFinal/_temp_division_table.docx')
    print('Division table saved.')


if __name__ == '__main__':
    build_process_report()
    build_division_table()
    print('All done.')
